"""
Ingestion pipeline using LlamaIndex components.

Flow per file:
  text extraction
  → LlamaIndex Document (doc_id = file_id for ref tracking)
  → IngestionPipeline: parse → extract metadata → embed
  → quality filter
  → MongoDocumentStore.add_documents()   ← full metadata, all nodes
  → strip large metadata fields from leaf nodes
  → PineconeVectorStore.add(leaf_nodes)  ← vectors + minimal metadata
  → files status update
"""
import asyncio
import logging
from pathlib import Path
from typing import Any

from bson import ObjectId
from llama_index.core.extractors import (
    QuestionsAnsweredExtractor,
    SummaryExtractor,
    TitleExtractor,
)
from llama_index.core.ingestion import IngestionPipeline
from llama_index.core.node_parser import (
    CodeSplitter,
    HierarchicalNodeParser,
    JSONNodeParser,
    MarkdownNodeParser,
    SentenceSplitter,
    get_leaf_nodes,
)
from llama_index.core.schema import BaseNode, Document, MetadataMode, NodeRelationship, TextNode
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.llms.openai import OpenAI
from llama_index.storage.docstore.mongodb import MongoDocumentStore
from llama_index.vector_stores.pinecone import PineconeVectorStore

from app.cache import invalidate_bm25_cache
from app.config import get_settings
from app.database import get_database
from app.models.file import FileStatus
from app.pinecone_client import get_index
from app.repositories import FileRepository, FolderRepository

logger = logging.getLogger(__name__)

MIN_CHUNK_CHARS = 50
EMBED_BATCH_SIZE = 100

CODE_EXTENSIONS = {".py", ".js", ".ts", ".go", ".java", ".cpp", ".c", ".cs", ".rb", ".rs", ".swift"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tiff"}

_CODE_LANGUAGE_MAP: dict[str, str] = {
    ".py": "python", ".js": "javascript", ".ts": "typescript",
    ".go": "go", ".java": "java", ".cpp": "cpp", ".c": "c",
    ".cs": "c_sharp", ".rb": "ruby", ".rs": "rust", ".swift": "swift",
}

# Extractor-generated fields that are large strings. They enrich the embedding
# at compute time but must be stripped before storing as Pinecone vector metadata
# (Pinecone limit: 40 KB per vector). Full values are kept in the docstore.
_LARGE_METADATA_KEYS = {
    "questions_this_excerpt_can_answer",
    "section_summary",
    "prev_section_summary",
    "next_section_summary",
    "excerpt_keywords",
}


# ---------------------------------------------------------------------------
# Shared docstore builder (same namespace used in ingestion + retrieval)
# ---------------------------------------------------------------------------

def _build_docstore(folder_id: str, settings: Any) -> MongoDocumentStore:
    db_name = settings.MONGODB_URI.rsplit("/", 1)[-1].split("?")[0] or "ragdb"
    return MongoDocumentStore.from_uri(
        uri=settings.MONGODB_URI,
        db_name=db_name,
        namespace=f"llamaindex_{folder_id}",
    )


# ---------------------------------------------------------------------------
# Status helpers
# ---------------------------------------------------------------------------

async def _set_status(file_oid: ObjectId, status: FileStatus, error: str | None = None) -> None:
    await FileRepository(get_database()).update_status(str(file_oid), status, error=error)


# ---------------------------------------------------------------------------
# Text extraction helpers
# ---------------------------------------------------------------------------

def _extract_pdf_text(path: str) -> str:
    try:
        import pypdf
        reader = pypdf.PdfReader(path)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        if text.strip():
            return text
    except Exception as exc:
        logger.warning("[INGESTION] pypdf failed for %s: %s — trying OCR", path, exc)
    return _ocr_file(path)


def _extract_docx_text(path: str) -> str:
    import docx
    doc = docx.Document(path)
    return "\n".join(p.text for p in doc.paragraphs)


def _ocr_file(path: str) -> str:
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(path)
        result = pytesseract.image_to_string(img)
        logger.debug("[INGESTION] OCR extracted %d chars from %s", len(result), path)
        return result
    except Exception as exc:
        logger.warning("[INGESTION] OCR failed for %s: %s", path, exc)
        return ""


async def _image_caption_via_openai(path: str, api_key: str) -> str:
    import base64
    from openai import AsyncOpenAI
    client = AsyncOpenAI(api_key=api_key)
    ext = Path(path).suffix.lstrip(".").lower()
    mime = f"image/{ext}" if ext in {"png", "jpg", "jpeg", "webp", "gif"} else "image/png"
    with open(path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode()
    response = await client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": [
            {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
            {"type": "text", "text": "Describe this image concisely for a document search index."},
        ]}],
        max_tokens=300,
    )
    caption = response.choices[0].message.content or ""
    logger.debug("[INGESTION] Vision caption: %d chars", len(caption))
    return caption


# ---------------------------------------------------------------------------
# Document builder — doc_id = file_id so docstore tracks nodes per file
# ---------------------------------------------------------------------------

def _make_document(text: str, file_name: str, content_type: str, storage_path: str, file_id: str) -> Document:
    return Document(
        text=text,
        doc_id=file_id,
        metadata={
            "file_name": file_name,
            "content_type": content_type,
            "storage_path": storage_path,
        },
        metadata_template="{key}: {value}",
        text_template="Metadata:\n{metadata_str}\n\nContent:\n{content}",
    )


# ---------------------------------------------------------------------------
# Parser + pipeline selection
# ---------------------------------------------------------------------------

def _make_extractors(llm: OpenAI, include_summary: bool = True) -> list:
    base = [
        TitleExtractor(llm=llm, nodes=5),
        QuestionsAnsweredExtractor(llm=llm, questions=3),
    ]
    if include_summary:
        base.append(SummaryExtractor(llm=llm, summaries=["self"]))
    return base


def _build_pipeline(
    ext: str,
    content_type: str,
    text_size: int,
    llm: OpenAI,
    embed_model: OpenAIEmbedding,
) -> tuple[IngestionPipeline | None, HierarchicalNodeParser | None, IngestionPipeline | None, str]:
    """
    Returns (flat_pipeline, hier_parser, leaf_pipeline, strategy_label).

    For flat strategies: flat_pipeline is set, hier_parser and leaf_pipeline are None.
    For hierarchical strategies: flat_pipeline is None, hier_parser and leaf_pipeline are set.

    HierarchicalNodeParser cannot be used inside IngestionPipeline — the pipeline
    collapses the multi-level node output via deduplication. Instead we call the
    parser directly, use get_leaf_nodes(), then run extractors + embedding on leaf
    nodes only via leaf_pipeline.
    """
    flat_extractors = _make_extractors(llm, include_summary=True)
    # SummaryExtractor only accepts TextNode — safe for leaf-only pipelines
    leaf_extractors = _make_extractors(llm, include_summary=True)

    if ext == ".json":
        return IngestionPipeline(
            transformations=[JSONNodeParser(), *flat_extractors, embed_model],
        ), None, None, "json"

    if content_type.startswith("code:"):
        language = _CODE_LANGUAGE_MAP.get(ext, "python")
        try:
            parser: Any = CodeSplitter(language=language, chunk_lines=40, chunk_lines_overlap=5, max_chars=4000)
        except Exception:
            logger.warning("[INGESTION] CodeSplitter unavailable for %s — using SentenceSplitter", language)
            parser = SentenceSplitter(chunk_size=512, chunk_overlap=50)
        return IngestionPipeline(
            transformations=[parser, *flat_extractors, embed_model],
        ), None, None, f"code:{language}"

    if ext in {".md", ".markdown"}:
        if text_size < 4_000:
            return IngestionPipeline(
                transformations=[MarkdownNodeParser(), *flat_extractors, embed_model],
            ), None, None, "markdown_only"
        if text_size < 16_000:
            return IngestionPipeline(
                transformations=[
                    MarkdownNodeParser(),
                    SentenceSplitter(chunk_size=512, chunk_overlap=50),
                    *flat_extractors,
                    embed_model,
                ],
            ), None, None, "markdown_sentence"
        hier_parser = HierarchicalNodeParser.from_defaults(chunk_sizes=[2048, 512])
        leaf_pipeline = IngestionPipeline(transformations=[*leaf_extractors, embed_model])
        return None, hier_parser, leaf_pipeline, "markdown_hierarchical"

    if text_size < 16_000:
        return IngestionPipeline(
            transformations=[SentenceSplitter(chunk_size=512, chunk_overlap=50), *flat_extractors, embed_model],
        ), None, None, "sentence"

    hier_parser = HierarchicalNodeParser.from_defaults(chunk_sizes=[2048, 512])
    leaf_pipeline = IngestionPipeline(transformations=[*leaf_extractors, embed_model])
    return None, hier_parser, leaf_pipeline, "hierarchical"


# ---------------------------------------------------------------------------
# Quality filter
# ---------------------------------------------------------------------------

def _filter_nodes(nodes: list[BaseNode]) -> tuple[list[BaseNode], int]:
    before = len(nodes)
    filtered = [n for n in nodes if len(n.get_content(metadata_mode=MetadataMode.NONE)) >= MIN_CHUNK_CHARS]
    return filtered, before - len(filtered)


# ---------------------------------------------------------------------------
# Main ingestion entry point
# ---------------------------------------------------------------------------

async def run_ingestion(file_id: str) -> None:
    settings = get_settings()
    db = get_database()
    file_oid = ObjectId(file_id)
    file_repo = FileRepository(db)
    folder_repo = FolderRepository(db)

    file_doc = await file_repo.get_by_id(file_id)
    if file_doc is None:
        logger.error("[INGESTION] File %s not found in DB", file_id)
        return

    folder_doc = await folder_repo.get_by_id(str(file_doc["folder_id"]))
    if folder_doc is None:
        logger.error("[INGESTION] Folder not found for file %s", file_id)
        return

    storage_path: str = file_doc["storage_path"]
    file_name: str = file_doc["file_name"]
    ext: str = file_doc["file_extension"].lower()
    content_type: str = file_doc["content_type"]
    folder_oid: ObjectId = file_doc["folder_id"]
    namespace: str = folder_doc["pinecone_namespace"]

    logger.info("[INGESTION] Starting  file_id=%s  name=%r  content_type=%s  namespace=%s",
                file_id, file_name, content_type, namespace)

    # --- Phase 1: Text extraction ---
    await _set_status(file_oid, FileStatus.PARSING)
    logger.info("[INGESTION] Phase=PARSING  file=%r  ext=%s", file_name, ext)
    try:
        if content_type == "markdown" or ext in {".md", ".markdown"}:
            text = Path(storage_path).read_text(errors="replace")
        elif content_type == "pdf":
            text = _extract_pdf_text(storage_path)
        elif content_type == "word":
            text = _extract_docx_text(storage_path)
        elif content_type == "image":
            text = _ocr_file(storage_path)
            if not text.strip():
                logger.info("[INGESTION] OCR empty — using OpenAI vision caption")
                text = await _image_caption_via_openai(storage_path, settings.OPENAI_API_KEY)
            ext = ".txt"
            content_type = "text"
        else:
            text = Path(storage_path).read_text(errors="replace")

        logger.info("[INGESTION] Extracted %d chars from %r", len(text), file_name)
    except Exception as exc:
        logger.exception("[INGESTION] Text extraction FAILED for file %s", file_id)
        await _set_status(file_oid, FileStatus.FAILED, str(exc))
        return

    # --- Phase 2: Build LlamaIndex objects ---
    llm = OpenAI(model="gpt-4o-mini", api_key=settings.OPENAI_API_KEY, temperature=0)
    embed_model = OpenAIEmbedding(
        model=settings.EMBEDDING_MODEL,
        api_key=settings.OPENAI_API_KEY,
        dimensions=settings.EMBEDDING_DIMENSIONS if settings.EMBEDDING_DIMENSIONS != 1536 else None,
        embed_batch_size=EMBED_BATCH_SIZE,
    )

    flat_pipeline, hier_parser, leaf_pipeline, parsing_strategy = _build_pipeline(
        ext=ext, content_type=content_type, text_size=len(text), llm=llm, embed_model=embed_model,
    )
    logger.info("[INGESTION] Pipeline strategy=%s", parsing_strategy)

    # --- Phase 3: parse → extract metadata → embed ---
    await _set_status(file_oid, FileStatus.EMBEDDING)
    logger.info("[INGESTION] Phase=PIPELINE_RUN  strategy=%s", parsing_strategy)
    document = _make_document(text, file_name, content_type, storage_path, file_id)
    try:
        if flat_pipeline is not None:
            # Flat strategies: single pipeline handles parsing + extraction + embedding
            nodes: list[BaseNode] = await flat_pipeline.arun(documents=[document], show_progress=False)
            logger.info("[INGESTION] Flat pipeline produced %d nodes", len(nodes))
        else:
            # Hierarchical strategies: HierarchicalNodeParser must run outside the pipeline
            # because IngestionPipeline deduplication collapses multi-level node output.
            # Pattern: parse all → get_leaf_nodes → run extractors + embedding on leaves only.
            assert hier_parser is not None and leaf_pipeline is not None
            all_parsed: list[BaseNode] = hier_parser.get_nodes_from_documents([document])
            leaf_raw = get_leaf_nodes(all_parsed)
            logger.info("[INGESTION] Hierarchical parser: %d total nodes, %d leaf nodes",
                        len(all_parsed), len(leaf_raw))
            # Run extractors + embedding on leaf nodes only
            embedded_leaves: list[BaseNode] = await leaf_pipeline.arun(nodes=leaf_raw, show_progress=False)
            # Combine: parent nodes (for docstore/AutoMerging) + embedded leaf nodes
            parent_ids = {n.node_id for n in embedded_leaves}
            parent_nodes = [n for n in all_parsed if n.node_id not in parent_ids]
            nodes = parent_nodes + embedded_leaves
            logger.info("[INGESTION] Hierarchical pipeline: %d parent + %d leaf = %d total nodes",
                        len(parent_nodes), len(embedded_leaves), len(nodes))
    except Exception as exc:
        logger.exception("[INGESTION] IngestionPipeline FAILED for file %s", file_id)
        await _set_status(file_oid, FileStatus.FAILED, str(exc))
        return

    if not nodes:
        logger.error("[INGESTION] Pipeline returned zero nodes for file %s", file_id)
        await _set_status(file_oid, FileStatus.FAILED, "Pipeline returned zero nodes")
        return

    # --- Phase 4: Quality filter (applied to all nodes; short parents are rare but drop them too) ---
    nodes, dropped = _filter_nodes(nodes)
    if dropped:
        logger.info("[INGESTION] Quality filter dropped %d short nodes", dropped)
    logger.info("[INGESTION] %d nodes after quality filter", len(nodes))

    # --- Phase 5: Save all nodes to MongoDocumentStore (full metadata preserved) ---
    # Docstore is the source of truth for node content, extracted metadata, and
    # parent/child relationships used by AutoMergingRetriever at query time.
    docstore = _build_docstore(str(folder_oid), settings)
    logger.info("[INGESTION] Saving %d nodes to MongoDocumentStore  namespace=llamaindex_%s",
                len(nodes), folder_oid)
    try:
        await asyncio.to_thread(docstore.add_documents, nodes, allow_update=True)
    except Exception as exc:
        logger.exception("[INGESTION] MongoDocumentStore save FAILED for file %s", file_id)
        await _set_status(file_oid, FileStatus.FAILED, str(exc))
        return

    # --- Phase 6: Strip large metadata fields, then upsert leaf nodes to Pinecone ---
    # The embeddings were already computed from the enriched text (with questions +
    # summaries prepended), so the vectors remain semantically rich. We only strip
    # the raw string fields from the Pinecone metadata dict to stay under 40 KB/vector.
    # Leaf nodes: TextNode with no CHILD relationships (parent nodes point to their children)
    leaf_nodes = [
        n for n in nodes
        if isinstance(n, TextNode) and NodeRelationship.CHILD not in n.relationships
    ]
    for node in leaf_nodes:
        for key in _LARGE_METADATA_KEYS:
            node.metadata.pop(key, None)

    logger.info("[INGESTION] Upserting %d leaf nodes to Pinecone  namespace=%s", len(leaf_nodes), namespace)
    try:
        pinecone_index = get_index()
        vector_store = PineconeVectorStore(pinecone_index=pinecone_index, namespace=namespace)
        await asyncio.to_thread(vector_store.add, leaf_nodes)
    except Exception as exc:
        logger.exception("[INGESTION] Pinecone upsert FAILED for file %s", file_id)
        await _set_status(file_oid, FileStatus.FAILED, str(exc))
        return

    leaf_count = len(leaf_nodes)

    # --- Phase 7: Update file record + folder stats ---
    await file_repo.update_status(
        file_id, FileStatus.READY, chunk_count=leaf_count, parsing_strategy=parsing_strategy
    )
    await folder_repo.increment_stats(str(folder_oid), file_delta=1, chunk_delta=leaf_count)

    # Invalidate BM25 cache so the next query rebuilds it with the new file's nodes
    await invalidate_bm25_cache(str(folder_oid))

    logger.info("[INGESTION] ✓ Complete  file_id=%s  name=%r  total_nodes=%d  leaf_nodes=%d  strategy=%s",
                file_id, file_name, len(nodes), leaf_count, parsing_strategy)


# ---------------------------------------------------------------------------
# Delete file data (called by DELETE /files/{id} and folder cascade delete)
# ---------------------------------------------------------------------------

async def delete_file_data(file_id: str, namespace: str) -> None:
    """
    Removes all nodes for a file from MongoDocumentStore and Pinecone.
    Uses ref_doc_id tracking that LlamaIndex sets automatically when doc_id
    is provided on the source Document (see _make_document).
    """
    settings = get_settings()
    db = get_database()
    file_repo = FileRepository(db)
    folder_repo = FolderRepository(db)

    # Fetch file before deletion so we can decrement folder stats
    file_doc = await file_repo.get_by_id(file_id)

    # Remove all nodes from docstore (tracks which nodes belong to this file via ref_doc_id)
    if file_doc:
        folder_id = str(file_doc["folder_id"])
        docstore = _build_docstore(folder_id, settings)
        try:
            await asyncio.to_thread(docstore.delete_ref_doc, file_id, raise_error=False)
            logger.info("[INGESTION] Removed nodes from docstore  file_id=%s", file_id)
        except Exception:
            logger.warning("[INGESTION] docstore.delete_ref_doc failed  file_id=%s", file_id)

    # Remove vectors from Pinecone
    try:
        pinecone_index = get_index()
        vector_store = PineconeVectorStore(pinecone_index=pinecone_index, namespace=namespace)
        await asyncio.to_thread(vector_store.delete, ref_doc_id=file_id)
        logger.info("[INGESTION] Deleted Pinecone vectors  file_id=%s  namespace=%s", file_id, namespace)
    except Exception:
        logger.warning("[INGESTION] Pinecone delete failed  file_id=%s", file_id)

    if file_doc:
        folder_id = str(file_doc["folder_id"])
        await folder_repo.increment_stats(folder_id, file_delta=-1, chunk_delta=-(file_doc.get("chunk_count") or 0))
        await invalidate_bm25_cache(folder_id)
